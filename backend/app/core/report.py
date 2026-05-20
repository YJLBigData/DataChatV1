"""DOCX 报告生成 — 使用可配置的提示词模板生成上市公司级别的经营分析报告。

流程：
  1. 用户在前端选择 report_template_id（默认 = 飞鹤上市报告标准商业分析报告）
  2. 取出模板的 prompt，连同 question / answer / plan / sql 喂给 LLM
  3. LLM 输出结构化 markdown（核心结论 / 关键指标 / 异常风险 / 管理建议 / 跟进问题）
  4. DOCX 排版：封面 + 摘要表 + LLM 报告正文 + 明细表 + 附录（口径/SQL）

LLM 失败时降级用 answer.narrative + highlights 兜底，永不抛出。
"""
from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.shared import Pt

logger = logging.getLogger("datachat.report")


def generate_report(
    question: str,
    answer: dict[str, Any],
    plan: dict[str, Any],
    sql: str,
    *,
    output_dir: str | Path,
    template_prompt: Optional[str] = None,
    template_name: str = "标准商业分析报告",
    llm=None,
) -> Path:
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    answer = answer or {}
    plan = plan or {}

    # ---- 1. LLM 生成正文（按模板提示词） ----
    body_markdown = _llm_compose_body(question, answer, plan, template_prompt, llm)

    # ---- 2. DOCX 排版 ----
    doc = Document()
    h = doc.add_heading("飞鹤经营分析报告", level=0)
    h.alignment = 1
    sub = doc.add_paragraph(template_name)
    sub.alignment = 1
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"业务问题：{question}")

    # 正文（LLM 输出）
    doc.add_heading("一、经营分析", level=1)
    for paragraph in body_markdown.split("\n\n"):
        text = paragraph.strip()
        if not text:
            continue
        # 简易 markdown：以 # 开头当二级标题，以 - 开头当列表
        if text.startswith("# "):
            doc.add_heading(text.lstrip("# ").strip(), level=2)
        elif text.startswith("## "):
            doc.add_heading(text.lstrip("# ").strip(), level=3)
        elif text.startswith("- "):
            for line in text.split("\n"):
                if line.startswith("- "):
                    doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(text)

    # 兜底：原始 narrative
    if answer.get("narrative") and not body_markdown:
        doc.add_paragraph(str(answer.get("narrative")))

    highlights = answer.get("highlights") or []
    if highlights:
        doc.add_heading("二、关键发现", level=1)
        for h in highlights:
            doc.add_paragraph(str(h), style="List Bullet")

    risks = answer.get("risk_notes") or []
    if risks:
        doc.add_heading("三、风险与提示", level=1)
        for r in risks:
            doc.add_paragraph(str(r), style="List Bullet")

    # 明细表
    table = answer.get("table") or {}
    columns = table.get("display_columns") or []
    rows = table.get("display_rows") or []
    if columns and rows:
        doc.add_heading("四、明细数据（前 50 行）", level=1)
        tbl = doc.add_table(rows=1, cols=len(columns))
        tbl.style = "Light Grid Accent 1"
        head_cells = tbl.rows[0].cells
        for i, c in enumerate(columns):
            head_cells[i].text = c.get("label") or c.get("key") or f"col{i}"
        for row in rows[:50]:
            cells = tbl.add_row().cells
            for i, value in enumerate(row[: len(columns)]):
                cells[i].text = str(value)

    # 附录
    doc.add_heading("附录 · 数据口径", level=1)
    metric_def = (answer.get("explainability") or {}).get("metric_definition") or {}
    if metric_def:
        doc.add_paragraph(f"指标：{metric_def.get('label','')}（{metric_def.get('name','')}）")
        doc.add_paragraph(f"表达式：{metric_def.get('expression','')}")
        doc.add_paragraph(f"来源表：{metric_def.get('table','')}")
        if metric_def.get("description"):
            doc.add_paragraph(f"说明：{metric_def['description']}")
    doc.add_paragraph(f"分组：{plan.get('group_by') or '（无）'}")
    doc.add_paragraph(f"过滤：{plan.get('filters') or '（无）'}")
    doc.add_paragraph(f"时间范围：{plan.get('time_range') or '（无）'}")
    if plan.get("calculation"):
        doc.add_paragraph(f"计算：{plan['calculation']}")

    doc.add_heading("附录 · SQL", level=1)
    p = doc.add_paragraph(sql or "")
    if p.runs:
        p.runs[0].font.name = "Menlo"
        p.runs[0].font.size = Pt(9)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = out_dir / f"feihe_report_{timestamp}.docx"
    doc.save(out_path)
    return out_path


def _llm_compose_body(
    question: str,
    answer: dict[str, Any],
    plan: dict[str, Any],
    template_prompt: Optional[str],
    llm,
) -> str:
    """用 LLM 按模板提示词生成专业正文。失败时降级。"""
    if not template_prompt:
        # 没传模板，直接用 narrative + highlights 拼装
        return _fallback_body(answer)
    if llm is None:
        try:
            from app.core.llm import get_llm_router
            llm = get_llm_router()
        except Exception:
            return _fallback_body(answer)

    table = answer.get("table") or {}
    rows_preview = (table.get("display_rows") or [])[:20]
    cols_label = [c.get("label") for c in (table.get("display_columns") or [])]
    table_text = " | ".join(cols_label) + "\n" + "\n".join(" | ".join(map(str, r)) for r in rows_preview)

    sys = template_prompt
    user = (
        f"业务问题：\n{question}\n\n"
        f"问数摘要：\n{answer.get('narrative') or ''}\n\n"
        f"关键发现：\n" + "\n".join(f"- {h}" for h in (answer.get('highlights') or [])) + "\n\n"
        f"风险提示：\n" + "\n".join(f"- {r}" for r in (answer.get('risk_notes') or [])) + "\n\n"
        f"数据明细（前 20 行）：\n{table_text}\n\n"
        "请按上市公司经营分析报告的格式输出。第一行写一句话核心结论；"
        "之后用 ##/## 二级三级标题分块；列表用 - 开头；不要展示 SQL；"
        "不要使用 markdown 代码块。"
    )
    try:
        res = llm.chat([{"role": "system", "content": sys}, {"role": "user", "content": user}],
                       temperature=0.0, max_tokens=2000)
        text = (res.text or "").strip()
        return text if text else _fallback_body(answer)
    except Exception as exc:
        logger.warning("report llm compose failed: %s — fallback", exc)
        return _fallback_body(answer)


def _fallback_body(answer: dict[str, Any]) -> str:
    parts: list[str] = []
    if answer.get("narrative"):
        parts.append(str(answer["narrative"]))
    if answer.get("highlights"):
        parts.append("## 关键发现\n" + "\n".join(f"- {h}" for h in answer["highlights"]))
    if answer.get("risk_notes"):
        parts.append("## 风险提示\n" + "\n".join(f"- {r}" for r in answer["risk_notes"]))
    return "\n\n".join(parts)
